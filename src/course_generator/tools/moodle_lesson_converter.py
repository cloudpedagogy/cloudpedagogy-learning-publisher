#!/usr/bin/env python3
"""
Moodle Lesson Word/HTML Export Converter
=======================================

Purpose
-------
Convert Moodle Lesson exports that look like Word `.doc` files but are actually
HTML into clean, editable, lesson-level Word documents that can be used as
canonical source files for your Quarto course publishing workflow.

Main design
-----------
This script is designed for your Word-first Quarto course engine.

It takes raw Moodle Lesson exports such as:

    imports/courses/intro_to_r/lesson_exports/Introduction to R 1. R and RStudio.doc

and creates clean authoring documents such as:

    imports/courses/intro_to_r/authoring_docx/01_r_and_rstudio.docx

The generated authoring DOCX files are intended to be edited by course teams
and then referenced from config/intro_to_r.yml.

Important architecture
----------------------
This converter does NOT generate final Quarto `.qmd` course pages.

Instead:

    Moodle Lesson export
        ↓
    moodle_lesson_converter.py
        ↓
    authoring_docx/     canonical editable Word source
    md/                 optional QA/debug Markdown
        ↓
    course_generator import-word
        ↓
    build/courses/.../*.qmd
        ↓
    Quarto render
        ↓
    output/courses/...

Outputs
-------
    authoring_docx/                  Main editable Word source docs
    page_docx/                       Optional QA page-level DOCX files
    md/                              Optional QA page-level Markdown files
    assets/<lesson_slug>/images/     Extracted images
    course_structure.generated.yml   YAML scaffold pointing to authoring_docx
    conversion_report.md             Human-readable QA report

Install
-------
    python -m pip install beautifulsoup4 python-docx pyyaml lxml

Example
-------
    python src/course_generator/tools/moodle_lesson_converter.py \
      imports/courses/intro_to_r/lesson_exports \
      --out imports/courses/intro_to_r \
      --module-id INTRO_R \
      --module-title "An Introduction to R" \
      --expected-lessons 1,2,3,4,5,6
"""

from __future__ import annotations

import argparse
import base64
import mimetypes
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

try:
    from bs4 import BeautifulSoup, NavigableString, Tag
except ImportError as exc:
    raise SystemExit(
        "Missing dependency: beautifulsoup4. Install with: python -m pip install beautifulsoup4"
    ) from exc

try:
    from docx import Document
    from docx.shared import Inches
except ImportError as exc:
    raise SystemExit(
        "Missing dependency: python-docx. Install with: python -m pip install python-docx"
    ) from exc

try:
    import yaml
except ImportError as exc:
    raise SystemExit(
        "Missing dependency: pyyaml. Install with: python -m pip install pyyaml"
    ) from exc


MOODLE_JUNK_PATTERNS = [
    "warning: this file must be opened using word",
    "moodle url:",
    "moodle language:",
    "course id:",
    "author username:",
    "image handling:",
    "contains embedded images:",
    "institution:",
]

NAVIGATION_TEXTS = {
    "next",
    "previous",
    "back",
    "continue",
    "description/jump",
    "jump",
    "start again",
}

QUESTION_TYPE_CODES = {"TF", "MC", "MA", "SA", "NUM", "MAT"}


@dataclass
class ExtractedImage:
    original_src: str
    path: Path
    alt: str = ""


@dataclass
class LessonPage:
    id: str
    number: int
    title: str
    html: str
    text_preview: str = ""
    images: list[ExtractedImage] = field(default_factory=list)
    classification: str = "content"
    word_count: int = 0


@dataclass
class LessonExport:
    input_path: Path
    lesson_title: str
    slug: str
    pages: list[LessonPage]
    authoring_docx_path: Optional[Path] = None


def slugify(value: str, fallback: str = "item") -> str:
    value = (value or "").strip().lower()
    value = re.sub(r"&", " and ", value)
    value = re.sub(r"[^a-z0-9]+", "-", value)
    value = re.sub(r"-+", "-", value).strip("-")
    return value or fallback


def safe_filename(value: str, fallback: str = "item") -> str:
    return slugify(value, fallback=fallback).replace("-", "_")


def read_text_with_fallback(path: Path) -> str:
    data = path.read_bytes()
    for encoding in ("utf-8", "windows-1252", "iso-8859-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def normalise_space(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def word_count(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text or ""))


def remove_global_export_junk(soup: BeautifulSoup) -> None:
    """Remove Word/Moodle document-level clutter, not lesson content."""
    for selector in ["style", "xml", "script", "meta", "link"]:
        for node in soup.find_all(selector):
            node.decompose()

    for node in list(soup.find_all(["p", "div"])):
        text = normalise_space(node.get_text(" ")).lower()
        if any(pattern in text for pattern in MOODLE_JUNK_PATTERNS):
            node.decompose()


def strip_navigation_from_fragment(fragment: BeautifulSoup) -> None:
    """
    Remove Moodle navigation artefacts from an already extracted content fragment.
    """
    for a in list(fragment.find_all("a")):
        text = normalise_space(a.get_text(" ")).lower()
        href = (a.get("href") or "").strip()
        if text in NAVIGATION_TEXTS or href == "#":
            parent = a.parent
            a.decompose()
            if parent and isinstance(parent, Tag):
                if not normalise_space(parent.get_text(" ")) and not parent.find("img"):
                    parent.decompose()

    for tag in list(fragment.find_all(["tr", "td", "th", "p"])):
        text = normalise_space(tag.get_text(" ")).lower()
        if not text:
            continue
        words = {w for w in re.split(r"\s+", text) if w}
        if text in NAVIGATION_TEXTS or (words and words.issubset(NAVIGATION_TEXTS)):
            if not tag.find("img"):
                tag.decompose()

    for tag in list(fragment.find_all(["td", "p"])):
        text = normalise_space(tag.get_text(" "))
        if text == "LE":
            tag.decompose()


def extract_lesson_title(soup: BeautifulSoup, input_path: Path) -> str:
    title_node = soup.find(class_=re.compile(r"MsoTitle", re.I))
    if title_node and normalise_space(title_node.get_text(" ")):
        return normalise_space(title_node.get_text(" "))

    html_title = soup.find("title")
    if html_title and normalise_space(html_title.get_text(" ")):
        title = normalise_space(html_title.get_text(" "))
        if title.lower() not in {"fred", "untitled"}:
            return title

    return input_path.stem


def find_lesson_page_nodes(soup: BeautifulSoup) -> list[Tag]:
    pages = soup.find_all("div", class_=lambda c: c and "chapter" in str(c).split())
    if pages:
        return pages

    pages = soup.find_all("div", id=re.compile(r"^page\d+", re.I))
    if pages:
        return pages

    body = soup.body or soup
    wrappers: list[Tag] = []
    current = None
    for child in list(body.children):
        if isinstance(child, Tag) and child.name == "h1":
            current = soup.new_tag("div")
            current["class"] = "chapter"
            wrappers.append(current)
        if current is not None:
            current.append(child.extract())

    return wrappers or [body]


def page_title(page_node: Tag, number: int) -> str:
    h = page_node.find(["h1", "h2", "h3"])
    if h and normalise_space(h.get_text(" ")):
        return normalise_space(h.get_text(" "))
    return f"Page {number}"


def extract_content_fragment(page_node: Tag) -> BeautifulSoup:
    fragments = page_node.find_all(
        "div", class_=lambda c: c and "no-overflow" in str(c).split()
    )

    if fragments:
        combined = BeautifulSoup("", "html.parser")
        root = combined.new_tag("div")
        combined.append(root)
        for frag in fragments:
            for child in list(frag.children):
                root.append(BeautifulSoup(str(child), "html.parser"))
        strip_navigation_from_fragment(combined)
        return combined

    clone = BeautifulSoup(str(page_node), "html.parser")
    for h in clone.find_all(["h1", "h2", "h3"], limit=1):
        h.decompose()
    strip_navigation_from_fragment(clone)
    return clone


def extract_data_image(src: str, asset_dir: Path, stem: str, index: int) -> Optional[Path]:
    match = re.match(r"data:(image/[a-zA-Z0-9.+-]+);base64,(.*)", src, flags=re.S)
    if not match:
        return None

    mime_type, b64_data = match.groups()
    ext = mimetypes.guess_extension(mime_type) or ".png"
    if ext == ".jpe":
        ext = ".jpg"

    path = asset_dir / f"{stem}_img_{index:03d}{ext}"
    path.parent.mkdir(parents=True, exist_ok=True)

    try:
        path.write_bytes(base64.b64decode(b64_data))
        return path
    except Exception:
        return None


def clean_page_html(
    page_node: Tag,
    asset_dir: Path,
    lesson_slug: str,
    page_num: int,
    out_dir: Path,
) -> tuple[str, list[ExtractedImage]]:
    fragment = extract_content_fragment(page_node)

    images: list[ExtractedImage] = []
    img_index = 1

    for img in fragment.find_all("img"):
        src = img.get("src", "")
        alt = img.get("alt", "") or ""

        if src.startswith("data:image/"):
            saved = extract_data_image(src, asset_dir, f"{lesson_slug}_p{page_num:03d}", img_index)
            if saved:
                try:
                    rel_src = saved.relative_to(out_dir).as_posix()
                except Exception:
                    rel_src = saved.as_posix()
                img["src"] = rel_src
                images.append(ExtractedImage(original_src=src[:80] + "...", path=saved, alt=alt))
                img_index += 1
            else:
                img.replace_with(fragment.new_string("[Image could not be extracted]"))
        else:
            images.append(ExtractedImage(original_src=src, path=Path(src), alt=alt))

    return str(fragment), images


def classify_page(title: str, html_clean: str) -> tuple[str, int]:
    soup = BeautifulSoup(html_clean, "html.parser")
    text = normalise_space(soup.get_text(" "))
    wc = word_count(text)
    low_title = title.strip().lower()
    low_text = text.lower()

    if low_title in {"end of course", "end"}:
        return "navigation_or_end_page", wc

    if low_title.startswith("page ") and wc <= 5 and not soup.find("img"):
        return "placeholder_or_branch_page", wc

    if "default mark" in low_text or "answers feedback grade" in low_text:
        return "quiz_or_question", wc

    tokens = set(low_text.split())
    if tokens.intersection({code.lower() for code in QUESTION_TYPE_CODES}) and (
        "grade" in low_text or "feedback" in low_text
    ):
        return "quiz_or_question", wc

    if wc < 10 and not soup.find("img"):
        return "sparse_review", wc

    return "content", wc


def parse_lesson(input_path: Path, out_dir: Path) -> LessonExport:
    html = read_text_with_fallback(input_path)
    soup = BeautifulSoup(html, "html.parser")

    lesson_title = extract_lesson_title(soup, input_path)
    lesson_slug = slugify(lesson_title, fallback=slugify(input_path.stem))

    remove_global_export_junk(soup)

    page_nodes = find_lesson_page_nodes(soup)
    pages: list[LessonPage] = []
    asset_dir = out_dir / "assets" / lesson_slug / "images"

    for idx, node in enumerate(page_nodes, start=1):
        title = page_title(node, idx)
        html_clean, images = clean_page_html(node, asset_dir, lesson_slug, idx, out_dir)
        text_preview = normalise_space(BeautifulSoup(html_clean, "html.parser").get_text(" "))[:250]

        if not text_preview and not images:
            continue

        classification, wc = classify_page(title, html_clean)
        pages.append(
            LessonPage(
                id=f"{lesson_slug}-p{idx:03d}",
                number=idx,
                title=title,
                html=html_clean,
                text_preview=text_preview,
                images=images,
                classification=classification,
                word_count=wc,
            )
        )

    return LessonExport(
        input_path=input_path,
        lesson_title=lesson_title,
        slug=lesson_slug,
        pages=pages,
    )


def add_run_text(paragraph, text: str) -> None:
    paragraph.add_run(text)


def add_inline_content_to_paragraph(paragraph, node: Tag | NavigableString) -> None:
    if isinstance(node, NavigableString):
        add_run_text(paragraph, str(node))
        return

    if not isinstance(node, Tag):
        return

    if node.name in {"strong", "b"}:
        run = paragraph.add_run(node.get_text())
        run.bold = True
    elif node.name in {"em", "i"}:
        run = paragraph.add_run(node.get_text())
        run.italic = True
    elif node.name == "a":
        text = normalise_space(node.get_text(" "))
        href = node.get("href")
        if href and href != "#":
            paragraph.add_run(f"{text} ({href})")
        else:
            paragraph.add_run(text)
    else:
        for child in node.children:
            add_inline_content_to_paragraph(paragraph, child)


def add_paragraph_from_tag(doc: Document, tag: Tag, style: Optional[str] = None) -> None:
    text = normalise_space(tag.get_text(" "))
    if not text:
        return

    try:
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    except Exception:
        p = doc.add_paragraph()

    if any(isinstance(c, Tag) and c.name in {"strong", "b", "em", "i", "a"} for c in tag.children):
        for child in tag.children:
            add_inline_content_to_paragraph(p, child)
    else:
        p.add_run(text)


def add_image_to_docx(doc: Document, img: Tag, base_dir: Path) -> None:
    src = img.get("src", "")
    alt = img.get("alt", "") or ""

    img_path = Path(src)
    if not img_path.is_absolute():
        img_path = base_dir / img_path

    if img_path.exists():
        try:
            doc.add_picture(str(img_path), width=Inches(5.8))
        except Exception:
            doc.add_paragraph(f"[Image could not be inserted: {src}]")
    else:
        doc.add_paragraph(f"[Image: {src}]")

    if alt:
        doc.add_paragraph(f"Alt text: {alt}")


def add_table_to_docx(doc: Document, table_tag: Tag) -> None:
    rows = []
    for row in table_tag.find_all("tr"):
        cells = [normalise_space(c.get_text(" ")) for c in row.find_all(["th", "td"])]
        low = " ".join(cells).strip().lower()
        if low in NAVIGATION_TEXTS or "description/jump" in low:
            continue
        if any(cells):
            rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    try:
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                table.cell(r_idx, c_idx).text = cell_text
    except Exception:
        for row in rows:
            doc.add_paragraph(" | ".join(row))


def add_html_to_docx(doc: Document, html: str, base_dir: Path) -> None:
    soup = BeautifulSoup(html, "html.parser")
    root = soup.find("div") or soup.body or soup

    def handle_node(node) -> None:
        if isinstance(node, NavigableString):
            text = normalise_space(str(node))
            if text:
                doc.add_paragraph(text)
            return

        if not isinstance(node, Tag):
            return

        if node.name == "h1":
            text = normalise_space(node.get_text(" "))
            if text:
                doc.add_heading(text, level=1)
        elif node.name in {"h2", "h3"}:
            text = normalise_space(node.get_text(" "))
            if text:
                doc.add_heading(text, level=2)
        elif node.name in {"h4", "h5", "h6"}:
            text = normalise_space(node.get_text(" "))
            if text:
                doc.add_heading(text, level=3)
        elif node.name == "p":
            imgs = node.find_all("img")
            text_without_imgs = normalise_space(node.get_text(" "))
            if text_without_imgs:
                add_paragraph_from_tag(doc, node)
            for img in imgs:
                add_image_to_docx(doc, img, base_dir)
        elif node.name in {"ul", "ol"}:
            for li in node.find_all("li", recursive=False):
                add_paragraph_from_tag(
                    doc,
                    li,
                    style="List Bullet" if node.name == "ul" else "List Number",
                )
        elif node.name == "li":
            add_paragraph_from_tag(doc, node, style="List Bullet")
        elif node.name == "img":
            add_image_to_docx(doc, node, base_dir)
        elif node.name == "table":
            add_table_to_docx(doc, node)
        elif node.name in {"div", "section", "article", "thead", "tbody"}:
            for child in node.children:
                handle_node(child)
        else:
            text = normalise_space(node.get_text(" "))
            if text:
                doc.add_paragraph(text)

    for child in root.children:
        handle_node(child)


def add_page_to_authoring_docx(
    doc: Document,
    page: LessonPage,
    out_dir: Path,
    include_review_notes: bool,
) -> None:
    doc.add_heading(page.title, level=1)

    if include_review_notes and page.classification in {
        "placeholder_or_branch_page",
        "navigation_or_end_page",
        "sparse_review",
    }:
        doc.add_paragraph(
            f"Review note :: imported Moodle page classified as {page.classification}; "
            f"word count {page.word_count}. Check whether this page should be kept, "
            f"merged, rewritten, or removed."
        )

    if page.classification == "quiz_or_question":
        doc.add_paragraph("Imported Moodle question")
        doc.add_paragraph(
            "Review note :: This Moodle question was preserved as editable text. "
            "Convert it manually to your course-engine quiz/activity syntax if required."
        )

    add_html_to_docx(doc, page.html, out_dir)


def html_to_markdown(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    root = soup.find("div") or soup.body or soup
    lines: list[str] = []

    def handle_node(node) -> None:
        if isinstance(node, NavigableString):
            text = normalise_space(str(node))
            if text:
                lines.append(text)
                lines.append("")
            return

        if not isinstance(node, Tag):
            return

        if node.name == "h1":
            lines.append(f"# {normalise_space(node.get_text(' '))}")
            lines.append("")
        elif node.name in {"h2", "h3"}:
            lines.append(f"## {normalise_space(node.get_text(' '))}")
            lines.append("")
        elif node.name in {"h4", "h5", "h6"}:
            lines.append(f"### {normalise_space(node.get_text(' '))}")
            lines.append("")
        elif node.name == "p":
            text = normalise_space(node.get_text(" "))
            if text:
                lines.append(text)
                lines.append("")
            for img in node.find_all("img"):
                handle_node(img)
        elif node.name == "ul":
            for li in node.find_all("li", recursive=False):
                text = normalise_space(li.get_text(" "))
                if text:
                    lines.append(f"- {text}")
            lines.append("")
        elif node.name == "ol":
            for i, li in enumerate(node.find_all("li", recursive=False), start=1):
                text = normalise_space(li.get_text(" "))
                if text:
                    lines.append(f"{i}. {text}")
            lines.append("")
        elif node.name == "li":
            text = normalise_space(node.get_text(" "))
            if text:
                lines.append(f"- {text}")
                lines.append("")
        elif node.name == "img":
            src = node.get("src", "")
            alt = node.get("alt", "") or ""
            lines.append(f"![{alt}]({src})")
            lines.append("")
        elif node.name == "table":
            rows = []
            for row in node.find_all("tr"):
                cells = [normalise_space(c.get_text(" ")) for c in row.find_all(["th", "td"])]
                low = " ".join(cells).strip().lower()
                if low in NAVIGATION_TEXTS or "description/jump" in low:
                    continue
                if any(cells):
                    rows.append(cells)

            if rows:
                max_cols = max(len(r) for r in rows)
                rows = [r + [""] * (max_cols - len(r)) for r in rows]
                lines.append("| " + " | ".join(rows[0]) + " |")
                lines.append("| " + " | ".join(["---"] * max_cols) + " |")
                for row in rows[1:]:
                    lines.append("| " + " | ".join(row) + " |")
                lines.append("")
        elif node.name in {"div", "section", "article", "thead", "tbody"}:
            for child in node.children:
                handle_node(child)
        else:
            text = normalise_space(node.get_text(" "))
            if text:
                lines.append(text)
                lines.append("")

    for child in root.children:
        handle_node(child)

    md = "\n".join(lines)
    md = re.sub(r"\n{3,}", "\n\n", md).strip()
    return md + "\n" if md else ""


def lesson_short_title(lesson_title: str) -> str:
    title = re.sub(r"^Introduction to R:\s*", "", lesson_title, flags=re.I)
    title = re.sub(r"^Introduction to R\s*\d+\.?\s*", "", title, flags=re.I)
    title = re.sub(r"^\d+\.\s*", "", title)
    return normalise_space(title) or lesson_title


def lesson_number_from_title_or_filename(
    lesson_title: str,
    input_path: Path,
    default_index: int,
) -> int:
    candidates = [lesson_title, input_path.stem]
    for c in candidates:
        m = re.search(r"(?:Introduction to R[:\s-]*)?(\d+)[\.\s-]", c, flags=re.I)
        if m:
            try:
                return int(m.group(1))
            except Exception:
                pass
    return default_index


def write_authoring_docx(
    lesson: LessonExport,
    out_dir: Path,
    authoring_dir: Path,
    lesson_number: int,
    include_review_notes: bool,
) -> Path:
    doc = Document()
    doc.add_heading(lesson.lesson_title, level=0)
    doc.add_paragraph(f"Source export :: {lesson.input_path.name}")
    doc.add_paragraph(
        "Import note :: This is a cleaned Moodle Lesson authoring document. "
        "Edit this file, not the original Moodle export."
    )

    for page in lesson.pages:
        doc.add_page_break()
        add_page_to_authoring_docx(doc, page, out_dir, include_review_notes)

    filename = (
        f"{lesson_number:02d}_"
        f"{safe_filename(lesson_short_title(lesson.lesson_title), fallback=lesson.slug)}.docx"
    )
    path = authoring_dir / filename
    doc.save(path)
    lesson.authoring_docx_path = path
    return path


def write_qa_outputs(lesson: LessonExport, out_dir: Path) -> None:
    """
    Write optional QA/debug artefacts.

    These are not the final Quarto course pages. Final QMD files are generated
    later by the main course engine during import-word.
    """
    page_docx_dir = out_dir / "page_docx"
    md_dir = out_dir / "md"

    page_docx_dir.mkdir(parents=True, exist_ok=True)
    md_dir.mkdir(parents=True, exist_ok=True)

    for page in lesson.pages:
        page_slug = slugify(page.title, fallback=f"page-{page.number:03d}")
        page_stem = f"{lesson.slug}_p{page.number:03d}_{page_slug}"

        page_doc = Document()
        add_page_to_authoring_docx(page_doc, page, out_dir, include_review_notes=True)
        page_doc.save(page_docx_dir / f"{page_stem}.docx")

        md_text = f"# {page.title}\n\n"
        md_text += html_to_markdown(page.html)
        (md_dir / f"{page_stem}.md").write_text(md_text, encoding="utf-8")


def build_course_yaml(
    exports: list[LessonExport],
    module_id: str,
    module_code: str,
    module_title: str,
    module_description: str,
) -> dict:
    yaml_sessions = []

    for idx, lesson in enumerate(exports, start=1):
        lesson_num = lesson_number_from_title_or_filename(lesson.lesson_title, lesson.input_path, idx)
        short_title = lesson_short_title(lesson.lesson_title)
        session_id = f"{module_id}-se{lesson_num:02d}"
        section_id = f"{session_id}-sec01"
        subpage_id = f"{section_id}-sp01"
        source_docx = lesson.authoring_docx_path.as_posix() if lesson.authoring_docx_path else ""

        yaml_sessions.append(
            {
                "id": session_id,
                "code": f"SE{lesson_num:02d}",
                "title": short_title,
                "type": "standard",
                "required": True,
                "source_status": "imported",
                "overview": f"Imported and cleaned from Moodle Lesson export: {lesson.input_path.name}",
                "learning_objectives": [],
                "render_mode": "multi_page",
                "sections": [
                    {
                        "id": section_id,
                        "title": short_title,
                        "kind": "section_overview",
                        "number": 1,
                        "navigation_style": "numbered_subpages",
                        "subpage_count": 1,
                        "subpages": [
                            {
                                "id": subpage_id,
                                "title": short_title,
                                "kind": "text_page",
                                "source_docx": source_docx,
                            }
                        ],
                    }
                ],
            }
        )

    return {
        "module": {
            "id": module_id,
            "code": module_code,
            "title": module_title,
            "description": module_description,
            "default_render_mode": "multi_page",
            "default_subpage_count": 1,
        },
        "sessions": yaml_sessions,
    }


def add_missing_lesson_placeholders(
    yaml_obj: dict,
    module_id: str,
    expected_lessons: list[int],
) -> None:
    sessions = yaml_obj.get("sessions", [])
    existing_nums = set()

    for s in sessions:
        m = re.search(r"SE(\d+)", s.get("code", ""))
        if m:
            existing_nums.add(int(m.group(1)))

    for n in expected_lessons:
        if n in existing_nums:
            continue

        title = "Operators, Objects, Data Types, and Functions in R" if n == 3 else f"Missing Lesson {n}"
        session_id = f"{module_id}-se{n:02d}"
        section_id = f"{session_id}-sec01"
        subpage_id = f"{section_id}-sp01"

        sessions.append(
            {
                "id": session_id,
                "code": f"SE{n:02d}",
                "title": title,
                "type": "standard",
                "required": True,
                "source_status": "missing_export",
                "overview": "Placeholder created because the Moodle Lesson export was not available.",
                "learning_objectives": [],
                "render_mode": "multi_page",
                "sections": [
                    {
                        "id": section_id,
                        "title": title,
                        "kind": "section_overview",
                        "number": 1,
                        "navigation_style": "numbered_subpages",
                        "subpage_count": 1,
                        "subpages": [
                            {
                                "id": subpage_id,
                                "title": "Source Not Yet Available",
                                "kind": "placeholder",
                            }
                        ],
                    }
                ],
            }
        )

    sessions.sort(key=lambda s: s.get("code", ""))


def write_report(
    exports: list[LessonExport],
    out_dir: Path,
    module_id: str,
    module_title: str,
) -> None:
    report_lines = [
        "# Moodle Lesson Conversion Report",
        "",
        f"Module: {module_title}",
        f"Module ID: {module_id}",
        "",
        "## Output interpretation",
        "",
        "- `authoring_docx/` contains the main editable Word source files for the course engine.",
        "- `page_docx/` contains page-level extraction artefacts for QA/debugging.",
        "- `md/` contains page-level Markdown QA/debug artefacts.",
        "- `assets/` contains extracted images and supporting media.",
        "",
        "## Architecture note",
        "",
        "This converter does not generate final Quarto `.qmd` course pages. "
        "Final QMD files should be generated by the main course engine during `import-word`.",
        "",
        "## Cleaning approach",
        "",
        "- Real Moodle lesson content is extracted from `div.no-overflow` blocks.",
        "- Moodle navigation wrappers such as `Description/Jump`, `Next`, and table layout borders are removed.",
        "- Images embedded as base64 data are extracted to `assets/` and inserted into the authoring DOCX.",
        "",
    ]

    for lesson in exports:
        report_lines.append(f"## {lesson.lesson_title}")
        report_lines.append(f"Source: `{lesson.input_path}`")
        if lesson.authoring_docx_path:
            report_lines.append(f"Authoring DOCX: `{lesson.authoring_docx_path}`")
        report_lines.append(f"Detected pages: {len(lesson.pages)}")
        report_lines.append("")
        report_lines.append("| Page | Title | Class | Words | Images | Preview |")
        report_lines.append("|---:|---|---|---:|---:|---|")
        for page in lesson.pages:
            preview = page.text_preview.replace("|", "\\|")
            report_lines.append(
                f"| {page.number:03d} | {page.title} | {page.classification} | "
                f"{page.word_count} | {len(page.images)} | {preview} |"
            )
        report_lines.append("")

    (out_dir / "conversion_report.md").write_text("\n".join(report_lines), encoding="utf-8")


def expand_input_paths(inputs: list[str]) -> list[Path]:
    paths: list[Path] = []

    for p in inputs:
        path = Path(p)
        if path.is_dir():
            candidates = []
            candidates.extend(sorted(path.glob("*.doc")))
            candidates.extend(sorted(path.glob("*.html")))
            candidates.extend(sorted(path.glob("*.htm")))
            paths.extend(candidates)
        else:
            paths.append(path)

    seen = set()
    unique = []
    for p in paths:
        if p.name.startswith("~$"):
            continue
        rp = p.resolve()
        if rp not in seen:
            unique.append(rp)
            seen.add(rp)

    return unique


def main(argv: Optional[list[str]] = None) -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Convert Moodle Lesson Word/HTML exports into clean editable "
            "authoring DOCX, optional Markdown QA artefacts, and YAML."
        )
    )
    parser.add_argument(
        "inputs",
        nargs="+",
        help="Input .doc/.html files or folders exported from Moodle Lesson",
    )
    parser.add_argument(
        "--out",
        default="converted_moodle_lessons",
        help="Output directory",
    )
    parser.add_argument(
        "--module-id",
        default="INTRO_R",
        help="Generated module ID",
    )
    parser.add_argument(
        "--module-code",
        default=None,
        help="Generated module code; defaults to module ID",
    )
    parser.add_argument(
        "--module-title",
        default="Imported Moodle Lesson Course",
        help="Generated module title",
    )
    parser.add_argument(
        "--module-description",
        default="Generated from Moodle Lesson Word/HTML exports.",
        help="Generated module description",
    )
    parser.add_argument(
        "--no-qa-outputs",
        action="store_true",
        help="Skip optional page_docx and md QA/debug outputs",
    )
    parser.add_argument(
        "--no-page-outputs",
        action="store_true",
        help="Legacy alias for --no-qa-outputs",
    )
    parser.add_argument(
        "--no-review-notes",
        action="store_true",
        help="Do not insert review notes into authoring DOCX for navigation/placeholder pages",
    )
    parser.add_argument(
        "--expected-lessons",
        default="",
        help="Comma-separated expected lesson numbers, e.g. 1,2,3,4,5,6",
    )
    args = parser.parse_args(argv)

    out_dir = Path(args.out).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    input_paths = expand_input_paths(args.inputs)

    if not input_paths:
        print("No input files found.", file=sys.stderr)
        return 2

    missing = [p for p in input_paths if not p.exists()]
    if missing:
        for p in missing:
            print(f"Missing input: {p}", file=sys.stderr)
        return 2

    exports: list[LessonExport] = []

    for input_path in input_paths:
        print(f"Converting: {input_path.name}")
        lesson = parse_lesson(input_path, out_dir)
        print(f"  Lesson title: {lesson.lesson_title}")
        print(f"  Pages detected: {len(lesson.pages)}")
        if lesson.pages:
            print(f"  First page preview: {lesson.pages[0].text_preview[:120]}")
        exports.append(lesson)

    exports.sort(
        key=lambda lesson: lesson_number_from_title_or_filename(
            lesson.lesson_title,
            lesson.input_path,
            999,
        )
    )

    authoring_dir = out_dir / "authoring_docx"
    authoring_dir.mkdir(parents=True, exist_ok=True)

    skip_qa_outputs = args.no_qa_outputs or args.no_page_outputs

    for idx, lesson in enumerate(exports, start=1):
        lesson_num = lesson_number_from_title_or_filename(
            lesson.lesson_title,
            lesson.input_path,
            idx,
        )
        write_authoring_docx(
            lesson,
            out_dir,
            authoring_dir,
            lesson_num,
            include_review_notes=not args.no_review_notes,
        )

        if not skip_qa_outputs:
            write_qa_outputs(lesson, out_dir)

    yaml_obj = build_course_yaml(
        exports=exports,
        module_id=args.module_id,
        module_code=args.module_code or args.module_id,
        module_title=args.module_title,
        module_description=args.module_description,
    )

    if args.expected_lessons.strip():
        expected = [int(item.strip()) for item in args.expected_lessons.split(",") if item.strip()]
        add_missing_lesson_placeholders(yaml_obj, args.module_id, expected)

    (out_dir / "course_structure.generated.yml").write_text(
        yaml.safe_dump(yaml_obj, sort_keys=False, allow_unicode=True),
        encoding="utf-8",
    )

    write_report(exports, out_dir, args.module_id, args.module_title)

    print(f"\nDone. Outputs written to: {out_dir}")
    print("Key files/folders:")
    print(f"  - {out_dir / 'course_structure.generated.yml'}")
    print(f"  - {out_dir / 'conversion_report.md'}")
    print(f"  - {out_dir / 'authoring_docx'}")
    if not skip_qa_outputs:
        print(f"  - {out_dir / 'page_docx'}")
        print(f"  - {out_dir / 'md'}")
    print(f"  - {out_dir / 'assets'}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())