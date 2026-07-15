#!/usr/bin/env python3
"""
Legacy WebCAL printable HTML → split editable DOCX source files.

Run from project root:

    .venv/bin/python src/course_generator/tools/legacy_html_to_docx.py imports/courses/ctm101_f01 --dry-run

Then:

    .venv/bin/python src/course_generator/tools/legacy_html_to_docx.py imports/courses/ctm101_f01
"""

from __future__ import annotations

import argparse
import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document


INTERFACE_IMAGE_NAMES = {
    "interaction_arrow.png",
    "blackcir.gif",
    "textbox.gif",
    "lshtm_logo_printable.png",
}

SKIP_TEXT = {"top", "skip navigation"}


@dataclass
class Section:
    number: int
    original_title: str
    title: str
    slug: str
    nodes: list[Tag]


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def slugify(text: str) -> str:
    text = clean_text(text).lower()
    text = re.sub(r"^\d+\s+", "", text)
    text = re.sub(r"[^a-z0-9]+", "_", text)
    text = re.sub(r"_+", "_", text)
    return text.strip("_") or "section"


def strip_number(title: str) -> str:
    return re.sub(r"^\d+\s+", "", clean_text(title)).strip()


def assert_safe_course_dir(course_dir: Path) -> None:
    resolved = course_dir.resolve()
    parts = resolved.parts

    if "imports" not in parts or "courses" not in parts:
        raise SystemExit(
            "Refusing to run. Course folder must be inside imports/courses/, e.g.\n"
            "  imports/courses/ctm101_f01"
        )

    raw_dir = resolved / "raw"
    if not raw_dir.exists():
        raise SystemExit(f"Missing raw folder: {raw_dir}")


def find_html_and_assets(raw_dir: Path) -> tuple[Path, Path | None]:
    html_files = sorted(raw_dir.glob("*.html"))

    if len(html_files) != 1:
        raise SystemExit(
            f"Expected exactly one .html file in {raw_dir}; found {len(html_files)}"
        )

    html_path = html_files[0]
    expected_assets = raw_dir / f"{html_path.stem}_files"

    if expected_assets.exists():
        return html_path, expected_assets

    asset_dirs = [p for p in raw_dir.iterdir() if p.is_dir() and p.name.endswith("_files")]

    if len(asset_dirs) == 1:
        return html_path, asset_dirs[0]

    return html_path, None


def split_sections(soup: BeautifulSoup) -> list[Section]:
    body = soup.body or soup
    sections: list[Section] = []

    current_title: str | None = None
    current_nodes: list[Tag] = []

    for node in body.children:
        if not isinstance(node, Tag):
            continue

        if node.name in {"header", "nav"}:
            continue

        if node.name == "h2":
            original_title = clean_text(node.get_text(" "))
            lowered = original_title.lower()

            if lowered in {"hints", "answers"}:
                break

            if current_title is not None:
                number = len(sections) + 1
                title = strip_number(current_title)
                sections.append(
                    Section(
                        number=number,
                        original_title=current_title,
                        title=title,
                        slug=slugify(title),
                        nodes=current_nodes,
                    )
                )

            current_title = original_title
            current_nodes = []
            continue

        if current_title is not None:
            current_nodes.append(node)

    if current_title is not None:
        number = len(sections) + 1
        title = strip_number(current_title)
        sections.append(
            Section(
                number=number,
                original_title=current_title,
                title=title,
                slug=slugify(title),
                nodes=current_nodes,
            )
        )

    return sections


def resolve_image_path(src: str, raw_dir: Path) -> Path | None:
    filename = Path(src.replace("\\", "/")).name
    if not filename:
        return None

    matches = list(raw_dir.rglob(filename))
    return matches[0] if matches else None


def copy_image(src_path: Path, assets_images_dir: Path, overwrite: bool) -> Path:
    assets_images_dir.mkdir(parents=True, exist_ok=True)
    dest = assets_images_dir / src_path.name

    if dest.exists() and not overwrite:
        return dest

    shutil.copy2(src_path, dest)
    return dest


def add_hyperlink_as_text(paragraph, text: str, href: str | None) -> None:
    text = clean_text(text)
    href = clean_text(href or "")

    if href and not href.startswith("file:"):
        paragraph.add_run(f"{text} ({href})")
    else:
        paragraph.add_run(text)


def add_inline_content(paragraph, node) -> None:
    if isinstance(node, NavigableString):
        paragraph.add_run(str(node))
        return

    if not isinstance(node, Tag):
        return

    if node.name in {"strong", "b"}:
        run = paragraph.add_run(clean_text(node.get_text(" ")))
        run.bold = True
        return

    if node.name in {"em", "i"}:
        run = paragraph.add_run(clean_text(node.get_text(" ")))
        run.italic = True
        return

    if node.name == "a":
        add_hyperlink_as_text(paragraph, node.get_text(" "), node.get("href"))
        return

    for child in node.children:
        add_inline_content(paragraph, child)


def add_paragraph_from_tag(doc: Document, tag: Tag, style: str | None = None) -> None:
    text = clean_text(tag.get_text(" "))
    if not text or text.lower() in SKIP_TEXT:
        return

    try:
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    except Exception:
        p = doc.add_paragraph()

    if any(isinstance(c, Tag) for c in tag.children):
        for child in tag.children:
            add_inline_content(p, child)
    else:
        p.add_run(text)


def add_table(doc: Document, table_tag: Tag) -> None:
    rows = table_tag.find_all("tr")
    if not rows:
        return

    grid = []
    for row in rows:
        cells = [clean_text(c.get_text(" ")) for c in row.find_all(["td", "th"])]
        if any(cells):
            grid.append(cells)

    if not grid:
        return

    max_cols = max(len(row) for row in grid)
    table = doc.add_table(rows=len(grid), cols=max_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(grid):
        for c_idx, cell_text in enumerate(row):
            table.cell(r_idx, c_idx).text = cell_text

    doc.add_paragraph("")


def add_image_block(
    doc: Document,
    img: Tag,
    raw_dir: Path,
    assets_images_dir: Path,
    report: list[str],
    overwrite: bool,
    caption: str = "",
) -> None:
    src = img.get("src", "")
    image_name = Path(src.replace("\\", "/")).name

    if image_name in INTERFACE_IMAGE_NAMES:
        report.append(f"- Ignored interface image: `{image_name}`")
        return

    src_path = resolve_image_path(src, raw_dir)
    alt = clean_text(img.get("alt", ""))

    if not src_path:
        report.append(f"- WARNING: image not found: `{src}`")
        return

    copied = copy_image(src_path, assets_images_dir, overwrite)
    rel = f"assets/images/{copied.name}"

    doc.add_paragraph(f"Image :: {rel}")
    doc.add_paragraph(f"Alt :: {alt or '[missing alt text — review]'}")

    if caption:
        doc.add_paragraph(f"Caption :: {caption}")

    doc.add_paragraph("Width :: 70%")
    doc.add_paragraph("")

    report.append(f"- Copied image: `{src_path.name}` → `{rel}`")

    if not alt:
        report.append(f"  - Missing alt text: `{src_path.name}`")


def add_html_node(
    doc: Document,
    node: Tag,
    raw_dir: Path,
    assets_images_dir: Path,
    report: list[str],
    overwrite: bool,
) -> None:
    if not isinstance(node, Tag):
        return

    if node.name == "p" and "pageid" in (node.get("class") or []):
        return

    text = clean_text(node.get_text(" "))
    if text.lower() in SKIP_TEXT:
        return

    if node.name in {"script", "style", "meta", "link"}:
        return

    if node.name in {"h3", "h4", "h5", "h6"}:
        level = 2 if node.name == "h3" else 3
        if text:
            doc.add_heading(text, level=level)
        return

    if node.name == "p" and "accordion_heading" in (node.get("class") or []):
        doc.add_heading(text, level=3)
        return

    if node.name == "p":
        imgs = node.find_all("img")
        if imgs:
            visible_text = text
            for ignored in [
                "Interaction",
                "radio button - choose one",
                "rectangle - text answer required",
            ]:
                visible_text = visible_text.replace(ignored, "")

            if clean_text(visible_text):
                add_paragraph_from_tag(doc, node)
            return

        add_paragraph_from_tag(doc, node)
        return

    if node.name in {"ul", "ol"}:
        style = "List Bullet" if node.name == "ul" else "List Number"
        for li in node.find_all("li", recursive=False):
            add_paragraph_from_tag(doc, li, style=style)
        return

    if node.name == "table":
        add_table(doc, node)
        return

    if node.name == "div":
        img = node.find("img", recursive=True)

        if img and "image_border" in (node.get("class") or []):
            caption_node = node.find(class_=re.compile("image_caption|image_credit"))
            caption = clean_text(caption_node.get_text(" ")) if caption_node else ""
            add_image_block(doc, img, raw_dir, assets_images_dir, report, overwrite, caption)
            return

        for child in node.children:
            if isinstance(child, Tag):
                add_html_node(doc, child, raw_dir, assets_images_dir, report, overwrite)
        return

    if node.name == "img":
        add_image_block(doc, node, raw_dir, assets_images_dir, report, overwrite)
        return

    if text:
        add_paragraph_from_tag(doc, node)


def write_section_docx(
    section: Section,
    course_dir: Path,
    raw_dir: Path,
    overwrite: bool,
    report: list[str],
) -> Path:
    docx_dir = course_dir / "docx"
    assets_images_dir = course_dir / "assets" / "images"

    docx_dir.mkdir(parents=True, exist_ok=True)
    assets_images_dir.mkdir(parents=True, exist_ok=True)

    filename = f"{section.number:02d}_{section.slug}.docx"
    output_path = docx_dir / filename

    if output_path.exists() and not overwrite:
        raise SystemExit(
            f"Refusing to overwrite existing file: {output_path}\n"
            "Use --overwrite if you want to replace generated files."
        )

    doc = Document()
    doc.add_heading(section.title, level=1)

    for node in section.nodes:
        add_html_node(doc, node, raw_dir, assets_images_dir, report, overwrite)

    doc.save(output_path)
    return output_path


def write_manifest(course_dir: Path, sections: list[Section], overwrite: bool) -> Path:
    manifest_path = course_dir / "manifest.yml"

    if manifest_path.exists() and not overwrite:
        raise SystemExit(
            f"Refusing to overwrite existing file: {manifest_path}\n"
            "Use --overwrite if you want to replace it."
        )

    lines = [
        "source: webcal_printable_html",
        "documents:",
    ]

    for section in sections:
        filename = f"{section.number:02d}_{section.slug}.docx"
        lines.append(f"  - number: {section.number}")
        lines.append(f'    title: "{section.title}"')
        lines.append(f'    source_docx: "docx/{filename}"')

    manifest_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return manifest_path


def write_report(
    course_dir: Path,
    html_path: Path,
    assets_dir: Path | None,
    sections: list[Section],
    created_docs: list[Path],
    report_lines: list[str],
    overwrite: bool,
) -> Path:
    report_path = course_dir / "conversion_report.md"

    if report_path.exists() and not overwrite:
        raise SystemExit(
            f"Refusing to overwrite existing file: {report_path}\n"
            "Use --overwrite if you want to replace it."
        )

    lines = [
        "# Legacy HTML to DOCX Conversion Report",
        "",
        f"Input HTML: `{html_path}`",
        f"Assets folder: `{assets_dir}`" if assets_dir else "Assets folder: not found",
        "",
        "## Sections detected",
        "",
    ]

    for section in sections:
        lines.append(f"- {section.number:02d}. {section.title}")

    lines.extend(["", "## Documents created", ""])

    for docx_path in created_docs:
        lines.append(f"- `{docx_path.relative_to(course_dir)}`")

    lines.extend(["", "## Image and QA notes", ""])

    if report_lines:
        lines.extend(report_lines)
    else:
        lines.append("- No image or QA notes recorded.")

    lines.extend(
        [
            "",
            "## Safety notes",
            "",
            "- The original `raw/` folder was not modified.",
            "- Generated Word files are in `docx/`.",
            "- Copied images are in `assets/images/`.",
            "- Interface images such as arrows, radio buttons and text boxes are ignored where detected.",
        ]
    )

    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report_path


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Split legacy WebCAL printable HTML into editable DOCX files."
    )
    parser.add_argument("course_dir", type=Path, help="Example: imports/courses/ctm101_f01")
    parser.add_argument("--dry-run", action="store_true", help="Preview without writing files.")
    parser.add_argument("--overwrite", action="store_true", help="Overwrite generated outputs.")

    args = parser.parse_args()

    course_dir = args.course_dir.resolve()
    assert_safe_course_dir(course_dir)

    raw_dir = course_dir / "raw"
    html_path, assets_dir = find_html_and_assets(raw_dir)

    html = html_path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "html.parser")
    sections = split_sections(soup)

    print("=" * 70)
    print("Legacy HTML to DOCX Importer")
    print("=" * 70)
    print(f"Course folder : {course_dir}")
    print(f"Raw folder    : {raw_dir}")
    print(f"HTML file     : {html_path.name}")
    print(f"Assets folder : {assets_dir.name if assets_dir else 'not found'}")
    print(f"Dry run       : {args.dry_run}")
    print(f"Overwrite     : {args.overwrite}")
    print()
    print(f"Detected sections: {len(sections)}")

    for section in sections:
        print(f"  {section.number:02d}. {section.title}")

    if args.dry_run:
        print()
        print("Dry run only. Nothing written.")
        print()
        print("Would create:")
        print(f"  {course_dir / 'docx'}")
        print(f"  {course_dir / 'assets' / 'images'}")
        print(f"  {course_dir / 'manifest.yml'}")
        print(f"  {course_dir / 'conversion_report.md'}")
        return

    report_lines: list[str] = []
    created_docs: list[Path] = []

    for section in sections:
        created_docs.append(
            write_section_docx(
                section=section,
                course_dir=course_dir,
                raw_dir=raw_dir,
                overwrite=args.overwrite,
                report=report_lines,
            )
        )

    manifest_path = write_manifest(course_dir, sections, overwrite=args.overwrite)
    report_path = write_report(
        course_dir=course_dir,
        html_path=html_path,
        assets_dir=assets_dir,
        sections=sections,
        created_docs=created_docs,
        report_lines=report_lines,
        overwrite=args.overwrite,
    )

    print()
    print("Created:")
    for path in created_docs:
        print(f"  {path.relative_to(course_dir)}")
    print(f"  {manifest_path.relative_to(course_dir)}")
    print(f"  {report_path.relative_to(course_dir)}")
    print()
    print("Done.")


if __name__ == "__main__":
    main()